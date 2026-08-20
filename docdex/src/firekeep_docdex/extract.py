"""Text extraction for the supported formats.

Three rules hold the whole module together:

* **It never raises.** A folder of documents is a folder of whatever a human
  happened to put there — a truncated PDF, a .docx that is really a renamed
  zip, a file that vanished between the walk and the read. Every one of those
  is a recorded per-file failure, never an exception that ends a sync over the
  other 4999 files.
* **An honest zero is a result, not a failure.** There is no OCR (a disclosed
  gap, I5): a scanned PDF yields no text. Reporting that as an error would put
  it in the retry set forever; reporting it as an empty extraction lets state
  record `seen_hash` and stop asking.
* **A file we decline to index is not a file that failed.** `.json` is only a
  document when it is a conversation export; a `package-lock.json` is noise,
  and calling that a failure would put it in the retry set and in the failure
  count forever. Those come back as an error string prefixed
  `unsupported: `, which the caller counts as skipped-unsupported and records
  as seen (I5 — the gap is said, not retried).

The HTML stripper is a `html.parser` subclass rather than a regex, and it is
deliberately the same construction maildex's mail stripper uses — copied, not
imported: neither wheel may depend on the other, and the shared thing is the
approach (nesting and character references handled by the parser, `<script>`
and `<style>` muted by a COUNTER rather than a flag), not the code.

**Provenance caveat for conversation exports.** A turn is labelled with its
role inside the extracted TEXT — `user:` / `assistant:` — and that is all.
The turns are not typed as claims, so nothing downstream can distinguish what
a person asserted from what a model generated: recall treats an exported chat
as one document, exactly like a PDF. Typed user-vs-model provenance is the
deferred Chatdex design (ROADMAP §5, 2026-08-19 evening amendment) and is
deliberately absent here rather than approximated.
"""
from __future__ import annotations

import json
from html.parser import HTMLParser
from pathlib import Path

from . import env_int

SUPPORTED_SUFFIXES = frozenset({
    ".md", ".txt", ".pdf", ".docx", ".html", ".htm", ".eml", ".json",
})

DEFAULT_MAX_EXTRACT_KB = 400

# The marker that separates "docdex declines to index this" from "docdex tried
# and failed". Both arrive as an error string, because extract() has exactly
# one channel for them; the prefix is what lets sync count them apart.
UNSUPPORTED_PREFIX = "unsupported: "

# Their content is markup, not prose. Without this, every HTML page ingests its
# own stylesheet and the recall snippet is a wall of CSS.
_SILENT = frozenset({"script", "style", "head", "title", "noscript"})

# Tags whose boundaries a reader perceives as a line break.
_BREAKS = frozenset({
    "p", "div", "br", "tr", "li", "h1", "h2", "h3", "h4", "h5", "h6",
    "table", "blockquote", "section", "article", "header", "footer", "pre", "hr",
})

# The mail headers a reader of a saved message wants, in the order they read
# them. Threading headers (Message-ID, In-Reply-To) are maildex's business: a
# .eml sitting in a documents folder is a document, not a mailbox.
_EML_HEADERS = (("Subject", "Subject"), ("From", "From"), ("To", "To"), ("Date", "Date"))

# Keys a chat export uses for the speaker and for what they said. Both families
# must be present for a dict to READ as a message — a bare `{"text": ...}` is
# any JSON object in the world.
_ROLE_KEYS = ("role", "author", "sender", "speaker", "from")
_TEXT_KEYS = ("content", "text", "message", "body", "parts")

# Where an object-wrapped export keeps its list of messages.
_MESSAGE_LIST_KEYS = ("messages", "chat_messages", "conversation", "turns", "history", "log")

# Guard on the recursive content walk. Export formats nest two or three deep;
# anything deeper is a hand-written file we have no business guessing about.
_MAX_CONTENT_DEPTH = 6


def is_supported(path: str | Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def is_unsupported(error: str | None) -> bool:
    """True when `error` is a decline rather than a failure."""
    return error is not None and error.startswith(UNSUPPORTED_PREFIX)


def max_extract_bytes() -> int:
    return env_int("FIREKEEP_DOCDEX_MAX_EXTRACT_KB", DEFAULT_MAX_EXTRACT_KB) * 1024


def truncate(text: str) -> tuple[str, bool]:
    """Cut `text` to the extract cap, returning `(text, truncated)`.

    The cap is a byte budget, so the cut is made on the encoded form and then
    decoded back with `errors="ignore"` — landing mid-codepoint would produce a
    JSON body the server cannot decode.
    """
    limit = max_extract_bytes()
    encoded = text.encode("utf-8")
    if len(encoded) <= limit:
        return text, False
    return encoded[:limit].decode("utf-8", "ignore"), True


def extract(path: str | Path) -> tuple[str, str | None]:
    """`(text, error)`. Exactly one of them is meaningful, except for the
    honest zero — `("", None)` — which is a valid, final result."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        return "", f"{UNSUPPORTED_PREFIX}file type '{p.suffix or p.name}'"
    try:
        if p.is_file() and p.stat().st_size == 0:
            # A zero-byte file is the honest zero, final, for EVERY format —
            # not a failure to retry. Empty just-created Office documents and
            # OneDrive folders produce them routinely, and handing an empty
            # stream to python-docx/pypdf raises package errors that read like
            # corruption: six empty .docx on the first real OneDrive source
            # sat as permanent "failures" WARN-ing in doctor forever.
            return "", None
        if suffix in (".md", ".txt"):
            return _text(p), None
        if suffix == ".pdf":
            return _pdf(p), None
        if suffix == ".docx":
            return _docx(p), None
        if suffix in (".html", ".htm"):
            return strip_html(_text(p)), None
        if suffix == ".eml":
            return _eml(p), None
        return _json(p)
    except Exception as e:  # noqa: BLE001 — per-file failures are DATA, never control flow
        return "", f"{type(e).__name__}: {e}"[:500]


def _text(p: Path) -> str:
    # errors="replace", not "strict": a stray Latin-1 byte in a notes file is a
    # document with one bad character, not an unreadable document.
    return p.read_bytes().decode("utf-8", "replace")


def _pdf(p: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — one unreadable page must not lose the rest
            continue
    return "\n\n".join(part for part in pages if part.strip())


def _docx(p: Path) -> str:
    import docx

    document = docx.Document(str(p))
    parts = [para.text for para in document.paragraphs]
    # Tables are ordinary document text — a runbook whose steps live in a table
    # would otherwise extract as an honest zero and never be indexed.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(part for part in parts if part.strip())


# --- html -------------------------------------------------------------------


class _Stripper(HTMLParser):
    """HTML in, readable text out. Never raises on malformed markup."""

    def __init__(self):
        # convert_charrefs=True is the default and is what turns `&amp;`,
        # `&#8217;` and `&nbsp;` into characters — including inside deeply
        # nested tags, which is the case a regex stripper gets wrong.
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        # A COUNTER, not a flag: `<style>` inside `<head>` would clear a flag
        # on its closing tag and leak the rest of the head into the output.
        self._muted = 0

    def handle_starttag(self, tag, attrs):
        if tag in _SILENT:
            self._muted += 1
        elif tag in _BREAKS:
            self._parts.append("\n")

    def handle_endtag(self, tag):
        if tag in _SILENT:
            self._muted = max(0, self._muted - 1)
        elif tag in _BREAKS:
            self._parts.append("\n")

    def handle_data(self, data):
        if not self._muted:
            self._parts.append(data)

    def text(self) -> str:
        return _collapse("".join(self._parts))


def strip_html(html: str) -> str:
    """Readable text from an HTML document. Never raises."""
    parser = _Stripper()
    try:
        parser.feed(html)
        parser.close()
    except Exception:  # noqa: BLE001 — malformed markup is DATA, not control flow
        pass
    return parser.text()


def _collapse(text: str) -> str:
    """Trim each line, drop runs of blank lines. Page markup emits enormous
    amounts of layout whitespace, and a corpus chunk that is 80% newlines
    wastes the embedding budget the real words needed."""
    lines = [line.strip() for line in
             text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    out: list[str] = []
    for line in lines:
        if line:
            out.append(line)
        elif out and out[-1]:
            out.append("")
    return "\n".join(out).strip()


# --- eml --------------------------------------------------------------------


def _eml(p: Path) -> str:
    """A saved message as a document: a small header block, then its body.

    Per-PART failures follow the `_pdf` precedent — a part with a charset that
    does not exist is skipped, and the rest of the message is still a document.
    Only a message we cannot parse at all raises, and the caller records that.
    """
    import email
    import email.policy

    parsed = email.message_from_bytes(p.read_bytes(), policy=email.policy.default)

    lines = []
    for name, label in _EML_HEADERS:
        value = _header(parsed, name)
        if value:
            lines.append(f"{label}: {value}")

    plain, html, attachments = _eml_parts(parsed)
    # text/plain wins when both are present: it is what the sender's client
    # wrote for a text reader, and it needs no stripping to be right.
    body = _collapse(plain) or strip_html(html)
    if attachments:
        lines.append(f"Attachments: {', '.join(attachments)}")
    if not lines:
        return body
    return "\n".join(lines) + ("\n\n" + body if body else "")


def _header(parsed, name: str) -> str:
    """One decoded header, never raising.

    A malformed `Subject` with a broken RFC 2047 encoded word raises from the
    policy's header parser on ACCESS, not on parse — so each one is read behind
    its own guard.
    """
    try:
        value = parsed.get(name)
        if value is None:
            return ""
        return " ".join(str(value).split())
    except Exception:  # noqa: BLE001
        return ""


def _eml_parts(parsed) -> tuple[str, str, list[str]]:
    """Plain text, HTML, and attachment NAMES — never attachment content.

    Attachment bodies are not extracted: a .eml carrying a 20MB PDF costs this
    function nothing, and the filename is the part a person searches for.
    """
    plain: list[str] = []
    html: list[str] = []
    attachments: list[str] = []

    for part in _walk_parts(parsed):
        try:
            disposition = (part.get_content_disposition() or "").lower()
            content_type = (part.get_content_type() or "").lower()
        except Exception:  # noqa: BLE001
            continue
        if disposition == "attachment" or content_type not in ("text/plain", "text/html"):
            name = _filename(part)
            if name:
                attachments.append(name)
            continue
        try:
            content = part.get_content()
        except Exception:  # noqa: BLE001 — a bad charset, a broken base64 block
            continue
        if isinstance(content, str):
            (plain if content_type == "text/plain" else html).append(content)

    return "\n\n".join(plain), "\n".join(html), attachments


def _walk_parts(parsed):
    try:
        if not parsed.is_multipart():
            return [parsed]
        return list(parsed.walk())
    except Exception:  # noqa: BLE001
        return [parsed]


def _filename(part) -> str:
    try:
        name = part.get_filename()
    except Exception:  # noqa: BLE001
        return ""
    if not name:
        return ""
    # A filename is sender-controlled text that lands in indexed content.
    # Newlines out, length bounded, and no path separators — `../../etc/passwd`
    # as a "filename" should read as nonsense, not as a path.
    cleaned = " ".join(str(name).split()).replace("/", "_").replace("\\", "_")
    return cleaned[:120]


# --- json: conversation exports only ----------------------------------------


def _json(p: Path) -> tuple[str, str | None]:
    """A `.json` file is a document only when it is a conversation export.

    Everything else is declined, not failed: an indexed `tsconfig.json` is
    noise in recall, and a generic dump of somebody's data model is worse.
    """
    try:
        data = json.loads(_text(p))
    except ValueError:
        return "", f"{UNSUPPORTED_PREFIX}not valid json"

    turns = _conversation(data)
    if turns is None:
        return "", f"{UNSUPPORTED_PREFIX}json is not a conversation export"
    if not turns:
        # Conversation-shaped, nothing said: an honest zero, recorded as seen
        # so an empty export is never re-extracted every cycle.
        return "", None

    body = "\n\n".join(f"{role}: {text}" for role, text in turns)
    return f"Conversation export ({len(turns)} turns)\n\n{body}", None


def _conversation(data) -> list[tuple[str, str]] | None:
    """The turns of a conversation export, or None when this is not one."""
    if isinstance(data, dict):
        return _conversation_object(data)
    if not isinstance(data, list):
        return None
    turns = _turns_from_list(data)
    if turns is not None:
        return turns
    # A ChatGPT `conversations.json` is a LIST of conversation objects, which
    # is the file a person actually exports — flattened into one document
    # because that is what the folder holds: one file, one document.
    collected: list[tuple[str, str]] = []
    matched = False
    for item in data:
        sub = _conversation_object(item)
        if sub is None:
            continue
        matched = True
        collected.extend(sub)
    return collected if matched else None


def _conversation_object(data) -> list[tuple[str, str]] | None:
    if not isinstance(data, dict):
        return None
    mapping = data.get("mapping")
    if isinstance(mapping, dict):
        return _chatgpt_turns(mapping)
    for key in _MESSAGE_LIST_KEYS:
        value = data.get(key)
        if isinstance(value, list):
            turns = _turns_from_list(value)
            if turns is not None:
                return turns
    return None


def _turns_from_list(items: list) -> list[tuple[str, str]] | None:
    """Turns from a flat `[{role, content}, ...]` list, or None.

    The majority test is what keeps a list of arbitrary records out: one object
    that happens to carry `author` and `text` does not make a chat log, and a
    real export is message dicts nearly all the way down.
    """
    if not items:
        return None
    shaped = [item for item in items if _looks_like_message(item)]
    if not shaped or len(shaped) * 2 < len(items):
        return None
    return [turn for turn in (_turn(item) for item in shaped) if turn is not None]


def _looks_like_message(item) -> bool:
    if not isinstance(item, dict):
        return False
    has_role = any(isinstance(item.get(key), (str, dict)) for key in _ROLE_KEYS)
    return has_role and any(key in item for key in _TEXT_KEYS)


def _chatgpt_turns(mapping: dict) -> list[tuple[str, str]] | None:
    """Turns from a ChatGPT export's `mapping` tree, in reading order.

    The mapping is a node tree, not a list: following `children` from the root
    is the only thing that reproduces the order a person saw. Nodes the walk
    never reaches (a pruned branch, a broken parent link) are appended after,
    so a damaged export loses its ordering rather than its content.
    """
    nodes = {key: value for key, value in mapping.items() if isinstance(value, dict)}
    if not any(isinstance(node.get("message"), dict) for node in nodes.values()):
        return None

    order: list[str] = []
    seen: set[str] = set()
    stack = [key for key, node in nodes.items() if not node.get("parent")]
    stack.reverse()
    while stack:
        key = stack.pop()
        if key in seen or key not in nodes:
            continue
        seen.add(key)
        order.append(key)
        children = nodes[key].get("children")
        if isinstance(children, list):
            stack.extend(reversed([c for c in children if isinstance(c, str)]))
    order.extend(key for key in nodes if key not in seen)

    turns = []
    for key in order:
        turn = _turn(nodes[key].get("message"))
        if turn is not None:
            turns.append(turn)
    return turns


def _turn(message) -> tuple[str, str] | None:
    if not isinstance(message, dict):
        return None
    text = ""
    for key in _TEXT_KEYS:
        if key in message:
            text = _content_text(message[key])
            if text.strip():
                break
    if not text.strip():
        return None
    return _role(message) or "unknown", _collapse(text)


def _role(message: dict) -> str:
    for key in _ROLE_KEYS:
        value = message.get(key)
        if isinstance(value, str) and value.strip():
            return " ".join(value.split())[:40]
        if isinstance(value, dict):
            for inner in ("role", "name"):
                nested = value.get(inner)
                if isinstance(nested, str) and nested.strip():
                    return " ".join(nested.split())[:40]
    return ""


def _content_text(value, depth: int = 0) -> str:
    """The text inside a `content` field, whatever shape the export chose.

    Covers the three in the wild: a bare string, ChatGPT's
    `{"parts": [...]}`, and Anthropic-style `[{"type": "text", "text": ...}]`.
    """
    if depth > _MAX_CONTENT_DEPTH:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        parts = [_content_text(item, depth + 1) for item in value]
        return "\n".join(part for part in parts if part.strip())
    if isinstance(value, dict):
        if isinstance(value.get("parts"), list):
            return _content_text(value["parts"], depth + 1)
        for key in ("text", "content", "value"):
            if key in value:
                return _content_text(value[key], depth + 1)
    return ""
